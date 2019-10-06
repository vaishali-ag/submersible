from tkinter import *
from tkinter import messagebox
from docx import Document
from docx.shared import Inches
import os
class create:
    def __init__(s):
    
        
        window.title("Wire Meter")
        l1=Label(window)
        l2=Label(window)
        l3=Label(window)
        
        labelSize=Label(window,text="Size Id:",width=18)
        labelOd=Label(window,text="O/D:")
        labelWeight=Label(window,text="Weight:")
        
        s.size=StringVar()
        s.od=StringVar()
        s.weight=StringVar()
        s.inSize=Entry(window,textvariable=s.size)
        s.inOd=Entry(window,textvariable=s.od)
        s.inWeight=Entry(window,textvariable=s.weight)

        b1=Button(window,text="Calculate" ,command=s.calc ,width=10)
        meter=s.calc
        b2=Button(window,text="Reset" ,command=s.reset ,width=10)
        b3=Button(window,text="Print" ,command=s.printing,width=10)
      

        s.labelLength1=Label(window)
        l1.grid(row=0)
        labelSize.grid(row=1,column=0)
        s.inSize.grid(row=1,column=1)

        l2.grid(row=2)

        labelOd.grid(row=3,column=0)
        s.inOd.grid(row=3,column=1)

        l3.grid(row=4)

        labelWeight.grid(row=5,column=0)
        s.inWeight.grid(row=5,column=1)
        s.labelLength1.grid(row=6)

        b1.grid(row=7,column=1)
        b2.grid(row=7,column=3)
        b3.grid(row=7 ,column =0)

    def reset(s):
        s.inSize.delete(0,END)
        s.inOd.delete(0,END)
        s.inWeight.delete(0,END)
    def printing(s):
        document = Document()
        size=s.size.get()
        od=(s.od.get())
        weight=(s.weight.get())
        
        m=s.calc()
        meter=str(m)
        document.add_heading('POLYSTER COATED', 0)

        p = document.add_paragraph('Winding Wire for Submersible Motor Pump')
        table = document.add_table(rows=2, cols=4)
        table.style='Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'SIZE I.D.'
        hdr_cells[0].border=1
        hdr_cells[1].text = size
        hdr_cells[2].text = 'O/D'
        hdr_cells[3].text = od
        row_cells = table.rows[1].cells
        row_cells[0].text='LENGTH'
        row_cells[1].text=meter
        row_cells[2].text='WEIGHT'
        row_cells[3].text=weight
        document.add_page_break()

        document.save('demo.docx')
        os.system('start demo.docx')


        
    def calc(s):

        
        size=s.size.get()
        od=int(s.od.get())
        weight=int(s.weight.get())
        
        s=['0-5','0-6','0-7','0-8','0-9','1-0','1-1','1-2','1-3','1-4','1-5','1-8','1-9','2-0','2-1','2-2','2-3','2-4']
        o=[125,135,145,155,165,175,130,140,150,160,170,180,190,200,210,220,230,240,80,90,100,110,120,130,140,150,160,170]
        if size in s and od in o:
            if size=="0-9" and od==125:
                meter=weight//6.200
            elif size=="1-0" and od==135:
                meter=weight//7.700
            elif size=="1-1" and od==145:
                meter=weight//9.100
            elif size=="1-2" and od==155:
                meter=weight//10.800
            elif size=="1-3" and od==165:
                meter=weight//12.400
            elif size=="1-4" and od==175:
                meter=weight//14.200
            
            elif size=="0-9" and od==130:
                meter=weight//6.400
            elif size=="1-0" and od==140:
                meter=weight//7.900
            elif size=="1-1" and od==150:
                meter=weight//9.400
            elif size=="1-2" and od==160:
                meter=weight//11.00
            elif size=="1-3" and od==170:
                meter=weight//12.800
            elif size=="1-4" and od==180:
                meter=weight//14.400
            elif size=="1-5" and od==190:
                meter=weight//16.000
            elif size=="1-6" and od==200:
                meter=weight//18.800
            elif size=="1-7" and od==210:
                meter=weight//22.400
            elif size=="1-8" and od==220:
                meter=weight//25.750
            elif size=="1-9" and od==230:
                meter=weight//29.100
            elif size=="2-0" and od==240:
                meter=weight//31.800
            elif size=="2-1" and od==250:
                meter=weight//34.500
            elif size=="2-2" and od==260:
                meter=weight//37.200

            
            elif size=="0-5" and od==80:
                meter=weight//2.000
            elif size=="0-6" and od==90:
                meter=weight//2.900
            elif size=="0-7" and od==100:
                meter=weight//3.900
            elif size=="0-8" and od==110:
                meter=weight//4.800
            elif size=="0-9" and od==120:
                meter=weight//6.000
            elif size=="1-0" and od==130:
                meter=weight//7.600
            elif size=="1-1" and od==140:
                meter=weight//9.000
            elif size=="1-2" and od==150:
                meter=weight//10.600
            elif size=="1-3" and od==160:
                meter=weight//12.200
            elif size=="1-4" and od==170:
                meter=weight//14.000
            messagebox.showinfo("Meter", meter)

           
        else:
            messagebox.showinfo("Error", "Invalid Input")
        return meter


window=Tk()
            
create()
window.mainloop()

